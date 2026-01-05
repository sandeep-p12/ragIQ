"""Template parser for PDF/DOCX templates without explicit placeholders."""

import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document as DocxDocument

from MemoIQ.config import MemoIQConfig
from MemoIQ.schema import FieldDefinition, TemplateSchema

logger = logging.getLogger(__name__)


class TemplateParser:
    """Parse templates to detect fillable fields."""
    
    def __init__(self, config: MemoIQConfig):
        """Initialize template parser."""
        self.config = config
    
    def parse_template(self, template_path: str) -> TemplateSchema:
        """
        Parse template file and detect fillable fields.
        
        Args:
            template_path: Path to template file (PDF or DOCX)
            
        Returns:
            TemplateSchema with detected fields
        """
        template_path_obj = Path(template_path)
        template_format = "pdf" if template_path_obj.suffix.lower() == ".pdf" else "docx"
        
        if template_format == "docx":
            return self._parse_docx_template(template_path)
        else:
            return self._parse_pdf_template(template_path)
    
    def _parse_docx_template(self, template_path: str) -> TemplateSchema:
        """Parse DOCX template."""
        docx_doc = DocxDocument(template_path)
        
        # Detect fillable fields
        field_definitions = {}
        structure_map = {
            "tables": [],
            "paragraphs": [],
            "formatting": {},
        }
        
        field_counter = 0
        
        # Process tables
        for table_idx, table in enumerate(docx_doc.tables):
            table_structure = {
                "table_id": f"table_{table_idx}",
                "rows": len(table.rows),
                "cols": len(table.columns) if table.rows else 0,
                "cells": [],
            }
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    
                    # Check if cell is empty or contains only whitespace
                    if not cell_text or cell_text.isspace():
                        # Empty cell - potential fillable field
                        field_id = f"field_{field_counter}"
                        field_counter += 1
                        
                        # Try to infer field name from context
                        context = self._get_cell_context(table, row_idx, col_idx)
                        field_name = self._infer_field_name(context, cell_text, "empty_cell")
                        
                        field_def = FieldDefinition(
                            field_id=field_id,
                            name=field_name,
                            type=self._infer_field_type(context, cell_text),
                            required=True,
                            description=f"Field in table {table_idx}, row {row_idx}, col {col_idx}",
                            location={
                                "type": "table_cell",
                                "table_id": f"table_{table_idx}",
                                "row": row_idx,
                                "col": col_idx,
                            },
                            field_type="empty_cell",
                            context=context,
                            placeholder_text="",
                        )
                        field_definitions[field_id] = field_def
                        
                        table_structure["cells"].append({
                            "row": row_idx,
                            "col": col_idx,
                            "field_id": field_id,
                            "is_empty": True,
                        })
                    else:
                        # Check for underscore patterns
                        if re.search(r'_+', cell_text):
                            # Contains underscores - might be a fillable field
                            field_id = f"field_{field_counter}"
                            field_counter += 1
                            
                            context = self._get_cell_context(table, row_idx, col_idx)
                            field_name = self._infer_field_name(context, cell_text, "underscore_line")
                            
                            field_def = FieldDefinition(
                                field_id=field_id,
                                name=field_name,
                                type=self._infer_field_type(context, cell_text),
                                required=True,
                                description=f"Field with underscores in table {table_idx}, row {row_idx}, col {col_idx}",
                                location={
                                    "type": "table_cell",
                                    "table_id": f"table_{table_idx}",
                                    "row": row_idx,
                                    "col": col_idx,
                                },
                                field_type="underscore_line",
                                context=context,
                                placeholder_text=cell_text,
                            )
                            field_definitions[field_id] = field_def
                            
                            table_structure["cells"].append({
                                "row": row_idx,
                                "col": col_idx,
                                "field_id": field_id,
                                "is_underscore": True,
                            })
                        else:
                            # Regular cell with content
                            table_structure["cells"].append({
                                "row": row_idx,
                                "col": col_idx,
                                "text": cell_text,
                                "is_content": True,
                            })
            
            structure_map["tables"].append(table_structure)
        
        # Process paragraphs for underscore lines and checkboxes
        for para_idx, para in enumerate(docx_doc.paragraphs):
            para_text = para.text.strip()
            
            if not para_text:
                continue
            
            # Check for underscore patterns
            if re.search(r'[A-Za-z\s]+:[\s]*_+', para_text):
                field_id = f"field_{field_counter}"
                field_counter += 1
                
                # Extract label and underscores
                match = re.match(r'([A-Za-z\s]+):[\s]*(_+)', para_text)
                if match:
                    label = match.group(1).strip()
                    underscores = match.group(2)
                    
                    field_def = FieldDefinition(
                        field_id=field_id,
                        name=label,
                        type=self._infer_field_type_from_label(label),
                        required=True,
                        description=f"Field with label '{label}'",
                        location={
                            "type": "paragraph",
                            "paragraph_idx": para_idx,
                            "run_idx": 0,
                        },
                        field_type="underscore_line",
                        context=label,
                        placeholder_text=underscores,
                    )
                    field_definitions[field_id] = field_def
                    
                    structure_map["paragraphs"].append({
                        "paragraph_idx": para_idx,
                        "field_id": field_id,
                        "is_underscore": True,
                    })
            
            # Check for checkboxes (☐ or similar)
            if '☐' in para_text or re.search(r'\[[\s]*\]', para_text):
                # Extract checkbox options
                checkbox_pattern = r'☐\s*([A-Za-z\s]+)'
                checkboxes = re.findall(checkbox_pattern, para_text)
                
                for checkbox_text in checkboxes:
                    field_id = f"field_{field_counter}"
                    field_counter += 1
                    
                    field_def = FieldDefinition(
                        field_id=field_id,
                        name=checkbox_text.strip(),
                        type="checkbox",
                        required=False,
                        description=f"Checkbox option: {checkbox_text.strip()}",
                        location={
                            "type": "paragraph",
                            "paragraph_idx": para_idx,
                            "checkbox_text": checkbox_text.strip(),
                        },
                        field_type="checkbox",
                        context=para_text,
                        placeholder_text=f"☐ {checkbox_text.strip()}",
                    )
                    field_definitions[field_id] = field_def
        
        # Extract sections
        sections = self._extract_sections(docx_doc)
        
        return TemplateSchema(
            template_path=template_path,
            template_format="docx",
            sections=sections,
            field_definitions=field_definitions,
            structure_map=structure_map,
        )
    
    def _parse_pdf_template(self, template_path: str) -> TemplateSchema:
        """Parse PDF template using ParseForge."""
        # Parse PDF to markdown first
        # STRICT: Use config.parsing_strategy if set, otherwise default to AUTO for templates
        from MemoIQ.rag.rag_adapter import parse_to_markdown
        from src.config.parsing_strategies import StrategyEnum
        
        # Use config strategy if set, otherwise AUTO for template structure detection
        strategy = self.config.parsing_strategy if self.config.parsing_strategy is not None else StrategyEnum.AUTO
        logger.info(f"Parsing TEMPLATE {Path(template_path).name} with strategy: {strategy}")
        if strategy == StrategyEnum.LLM_FULL:
            logger.warning("⚠️  Template parsing using LLM_FULL - structure detection may be limited")
        markdown = parse_to_markdown(template_path, self.config.parsing_config, strategy=strategy)
        
        # Parse markdown to detect structure
        # This is a simplified approach - in production, you'd want more sophisticated parsing
        field_definitions = {}
        structure_map = {
            "markdown": markdown,
            "tables": [],
            "sections": [],
        }
        
        field_counter = 0
        
        # Detect tables in markdown - improved parsing
        # Markdown tables have format: | col1 | col2 | col3 |
        # Empty cells appear as: | col1 |  | col3 |
        lines = markdown.split('\n')
        current_table = []
        table_idx = 0
        
        for line in lines:
            line = line.strip()
            # Check if this is a table row (starts and ends with |)
            if line.startswith('|') and line.endswith('|') and '|' in line[1:-1]:
                # This is a table row
                # Split by | and remove first/last empty elements
                cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove first/last empty from split
                current_table.append(cells)
            elif current_table:
                # End of table - process it
                if len(current_table) >= 2:  # At least header + 1 data row
                    # Process table rows (skip header row which is usually row 0)
                    for row_idx, row_cells in enumerate(current_table[1:], start=1):  # Start from row 1 (skip header)
                        for col_idx, cell in enumerate(row_cells):
                            # Check if cell is empty or contains only whitespace
                            if not cell or cell.isspace() or cell == '':
                                field_id = f"field_{field_counter}"
                                field_counter += 1
                                
                                # Try to infer field name from header or previous cell
                                context = ""
                                field_name = f"Table Field {field_counter}"
                                
                                # Get header if available
                                if current_table and col_idx < len(current_table[0]):
                                    header = current_table[0][col_idx].strip()
                                    if header:
                                        field_name = header
                                        context = f"Column header: {header}"
                                
                                # Get previous cell in same row for context
                                if col_idx > 0 and col_idx - 1 < len(row_cells):
                                    prev_cell = row_cells[col_idx - 1].strip()
                                    if prev_cell:
                                        context = f"{context} | Previous: {prev_cell}".strip()
                                        # Use previous cell as field name if no header
                                        if not field_name or field_name.startswith("Table Field"):
                                            field_name = prev_cell
                                
                                field_def = FieldDefinition(
                                    field_id=field_id,
                                    name=field_name,
                                    type=self._infer_field_type_from_label(field_name),
                                    required=True,
                                    description=f"Empty cell in table {table_idx}, row {row_idx}, col {col_idx}",
                                    location={
                                        "type": "markdown_table",
                                        "table_idx": table_idx,
                                        "row": row_idx,
                                        "col": col_idx,
                                    },
                                    field_type="empty_cell",
                                    context=context,
                                    placeholder_text="",
                                )
                                field_definitions[field_id] = field_def
                                logger.debug(f"Detected empty cell field: {field_id} = {field_name} (table {table_idx}, row {row_idx}, col {col_idx})")
                    
                    table_idx += 1
                current_table = []
        
        # Process last table if file ends with table
        if current_table and len(current_table) >= 2:
            for row_idx, row_cells in enumerate(current_table[1:], start=1):
                for col_idx, cell in enumerate(row_cells):
                    if not cell or cell.isspace() or cell == '':
                        field_id = f"field_{field_counter}"
                        field_counter += 1
                        
                        context = ""
                        field_name = f"Table Field {field_counter}"
                        if current_table and col_idx < len(current_table[0]):
                            header = current_table[0][col_idx].strip()
                            if header:
                                field_name = header
                                context = f"Column header: {header}"
                        
                        field_def = FieldDefinition(
                            field_id=field_id,
                            name=field_name,
                            type=self._infer_field_type_from_label(field_name),
                            required=True,
                            description=f"Empty cell in table {table_idx}, row {row_idx}, col {col_idx}",
                            location={
                                "type": "markdown_table",
                                "table_idx": table_idx,
                                "row": row_idx,
                                "col": col_idx,
                            },
                            field_type="empty_cell",
                            context=context,
                            placeholder_text="",
                        )
                        field_definitions[field_id] = field_def
        
        # Detect underscore lines
        underscore_pattern = r'([A-Za-z\s]+):[\s]*_+'
        for match in re.finditer(underscore_pattern, markdown):
            field_id = f"field_{field_counter}"
            field_counter += 1
            
            label = match.group(1).strip()
            
            field_def = FieldDefinition(
                field_id=field_id,
                name=label,
                type=self._infer_field_type_from_label(label),
                required=True,
                description=f"Field with label '{label}'",
                location={
                    "type": "markdown_text",
                    "match_start": match.start(),
                    "match_end": match.end(),
                },
                field_type="underscore_line",
                context=label,
                placeholder_text=match.group(0),
            )
            field_definitions[field_id] = field_def
        
        sections = self._extract_sections_from_markdown(markdown)
        
        return TemplateSchema(
            template_path=template_path,
            template_format="pdf",
            sections=sections,
            field_definitions=field_definitions,
            structure_map=structure_map,
        )
    
    def _get_cell_context(self, table, row_idx: int, col_idx: int) -> str:
        """Get context for a cell (nearby labels, headers)."""
        context_parts = []
        
        # Get row header (first cell in row)
        if row_idx < len(table.rows):
            first_cell = table.rows[row_idx].cells[0].text.strip()
            if first_cell and first_cell != table.rows[row_idx].cells[col_idx].text.strip():
                context_parts.append(f"Row header: {first_cell}")
        
        # Get column header (first row)
        if table.rows and col_idx < len(table.rows[0].cells):
            header_cell = table.rows[0].cells[col_idx].text.strip()
            if header_cell:
                context_parts.append(f"Column header: {header_cell}")
        
        # Get previous cell in same row
        if col_idx > 0 and row_idx < len(table.rows):
            prev_cell = table.rows[row_idx].cells[col_idx - 1].text.strip()
            if prev_cell:
                context_parts.append(f"Previous: {prev_cell}")
        
        return " | ".join(context_parts) if context_parts else ""
    
    def _infer_field_name(self, context: str, cell_text: str, field_type: str) -> str:
        """Infer field name from context."""
        # Try to extract meaningful name from context
        if "Row header:" in context:
            return context.split("Row header:")[1].split("|")[0].strip()
        elif "Column header:" in context:
            return context.split("Column header:")[1].split("|")[0].strip()
        elif "Previous:" in context:
            prev_text = context.split("Previous:")[1].split("|")[0].strip()
            # Remove common prefixes
            prev_text = re.sub(r'^(Name|Address|Amount|Date|Type|Status):?\s*', '', prev_text, flags=re.IGNORECASE)
            return prev_text or "Field"
        else:
            return "Field"
    
    def _infer_field_type(self, context: str, cell_text: str) -> str:
        """Infer field type from context and cell text."""
        context_lower = context.lower()
        cell_lower = cell_text.lower()
        
        if "amount" in context_lower or "price" in context_lower or "$" in cell_text:
            return "currency"
        elif "date" in context_lower:
            return "date"
        elif "number" in context_lower or "count" in context_lower:
            return "number"
        elif "email" in context_lower:
            return "text"  # Could be email, but using text for now
        else:
            return "text"
    
    def _infer_field_type_from_label(self, label: str) -> str:
        """Infer field type from label text."""
        label_lower = label.lower()
        
        if any(word in label_lower for word in ["amount", "price", "cost", "fee", "loan"]):
            return "currency"
        elif "date" in label_lower:
            return "date"
        elif any(word in label_lower for word in ["number", "count", "quantity"]):
            return "number"
        elif "email" in label_lower:
            return "text"
        else:
            return "text"
    
    def _extract_sections(self, docx_doc: DocxDocument) -> List[Dict[str, Any]]:
        """Extract sections from DOCX document."""
        sections = []
        current_section = None
        
        for para in docx_doc.paragraphs:
            style = para.style.name.lower()
            if "heading" in style or "title" in style:
                if current_section:
                    sections.append(current_section)
                
                level = int(style[-1]) if style[-1].isdigit() else 1
                current_section = {
                    "title": para.text.strip(),
                    "level": level,
                    "fields": [],
                }
            elif current_section and para.text.strip():
                # Add paragraph to current section
                pass
        
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _extract_sections_from_markdown(self, markdown: str) -> List[Dict[str, Any]]:
        """Extract sections from markdown."""
        sections = []
        heading_pattern = r'^(#{1,6})\s+(.+)$'
        
        for line in markdown.split('\n'):
            match = re.match(heading_pattern, line)
            if match:
                level = len(match.group(1))
                title = match.group(2).strip()
                sections.append({
                    "title": title,
                    "level": level,
                    "fields": [],
                })
        
        return sections

