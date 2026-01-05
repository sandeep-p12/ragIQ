"""Template filler - fills templates with extracted values while preserving structure."""

import logging
from pathlib import Path
from typing import Dict, List

from docx import Document as DocxDocument

from MemoIQ.schema import FieldExtraction, TemplateSchema

logger = logging.getLogger(__name__)


class TemplateFiller:
    """Fill templates with extracted values."""
    
    def __init__(self, config):
        """Initialize template filler."""
        self.config = config
    
    def fill_template(
        self,
        template_schema: TemplateSchema,
        extracted_fields: Dict[str, FieldExtraction],
        output_path: str,
    ) -> str:
        """
        Fill template with extracted values.
        
        Args:
            template_schema: Parsed template schema
            extracted_fields: Dict of field_id -> FieldExtraction
            output_path: Path to save filled DOCX
            
        Returns:
            Path to filled DOCX file
        """
        if template_schema.template_format == "docx":
            return self._fill_docx_template(template_schema, extracted_fields, output_path)
        else:
            # PDF: Parse to markdown, reconstruct as DOCX
            return self._fill_pdf_template(template_schema, extracted_fields, output_path)
    
    def _fill_docx_template(
        self,
        template_schema: TemplateSchema,
        extracted_fields: Dict[str, FieldExtraction],
        output_path: str,
    ) -> str:
        """Fill DOCX template."""
        # Load original template
        docx_doc = DocxDocument(template_schema.template_path)
        
        # Fill table cells
        filled_table_fields = 0
        for table_structure in template_schema.structure_map.get("tables", []):
            table_idx = int(table_structure["table_id"].split("_")[1])
            if table_idx < len(docx_doc.tables):
                table = docx_doc.tables[table_idx]
                
                for cell_info in table_structure.get("cells", []):
                    if "field_id" in cell_info:
                        field_id = cell_info["field_id"]
                        if field_id in extracted_fields:
                            field_extraction = extracted_fields[field_id]
                            row = cell_info["row"]
                            col = cell_info["col"]
                            
                            if row < len(table.rows) and col < len(table.rows[row].cells):
                                cell = table.rows[row].cells[col]
                                
                                # Clear cell and add value
                                cell.text = ""
                                value_str = str(field_extraction.value) if field_extraction.value is not None else ""
                                cell.add_paragraph(value_str)
                                
                                # Add citation if available
                                if field_extraction.citations:
                                    citation_text = self._format_citations(field_extraction.citations)
                                    if citation_text:
                                        # Add as footnote or in cell
                                        para = cell.paragraphs[0]
                                        run = para.add_run(f" [{citation_text}]")
                                        run.font.size = None  # Smaller font for citation
                                
                                filled_table_fields += 1
                                logger.debug(f"Filled table cell [{row}, {col}] with field '{field_id}': {value_str[:50]}")
        
        logger.info(f"Filled {filled_table_fields} table cells with extracted values")
        
        # Fill paragraphs with underscore lines
        filled_para_fields = 0
        for para_info in template_schema.structure_map.get("paragraphs", []):
            if "field_id" in para_info:
                field_id = para_info["field_id"]
                if field_id in extracted_fields:
                    field_extraction = extracted_fields[field_id]
                    para_idx = para_info["paragraph_idx"]
                    
                    if para_idx < len(docx_doc.paragraphs):
                        para = docx_doc.paragraphs[para_idx]
                        
                        # Replace underscore pattern with value
                        original_text = para.text
                        if "_" in original_text:
                            # Replace underscores with value
                            value_str = str(field_extraction.value) if field_extraction.value is not None else ""
                            new_text = original_text.replace("_" * 20, value_str)  # Replace long underscores
                            para.text = new_text
                            
                            # Add citation
                            if field_extraction.citations:
                                citation_text = self._format_citations(field_extraction.citations)
                                para.add_run(f" [{citation_text}]")
                            
                            filled_para_fields += 1
                            logger.debug(f"Filled paragraph {para_idx} with field '{field_id}': {value_str[:50]}")
        
        logger.info(f"Filled {filled_para_fields} paragraph fields with extracted values")
        
        # Update checkboxes
        for para in docx_doc.paragraphs:
            if '☐' in para.text:
                # Check if this paragraph has a field
                for field_id, field_extraction in extracted_fields.items():
                    if field_extraction.value is True and field_extraction.field_type == "checkbox":
                        # Replace ☐ with ☑
                        para.text = para.text.replace('☐', '☑', 1)
        
        # Save filled document
        docx_doc.save(output_path)
        logger.info(f"Filled template saved to {output_path}")
        
        return output_path
    
    def _fill_pdf_template(
        self,
        template_schema: TemplateSchema,
        extracted_fields: Dict[str, FieldExtraction],
        output_path: str,
    ) -> str:
        """Fill PDF template by converting to DOCX."""
        # Parse PDF to markdown
        # STRICT: Use config.parsing_strategy if set
        from MemoIQ.rag.rag_adapter import parse_to_markdown
        from src.config.parsing_strategies import StrategyEnum
        
        # Use config strategy if set, otherwise AUTO
        strategy = self.config.parsing_strategy if self.config.parsing_strategy is not None else StrategyEnum.AUTO
        logger.info(f"Parsing template for filling with strategy: {strategy}")
        markdown = parse_to_markdown(template_schema.template_path, self.config.parsing_config, strategy=strategy)
        
        # Replace fields in markdown
        filled_markdown = markdown
        for field_id, field_extraction in extracted_fields.items():
            if field_id in template_schema.field_definitions:
                field_def = template_schema.field_definitions[field_id]
                
                # Replace placeholder
                if field_def.placeholder_text:
                    value_str = str(field_extraction.value) if field_extraction.value is not None else ""
                    filled_markdown = filled_markdown.replace(field_def.placeholder_text, value_str)
                
                # Add citation
                if field_extraction.citations:
                    citation_text = self._format_citations(field_extraction.citations)
                    filled_markdown += f"\n[{field_def.name}: {citation_text}]\n"
        
        # Convert markdown to DOCX (simplified - would need proper markdown to DOCX converter)
        # For now, create a simple DOCX with the filled content
        docx_doc = DocxDocument()
        
        # Split markdown into paragraphs and add to DOCX
        for line in filled_markdown.split('\n'):
            if line.strip():
                para = docx_doc.add_paragraph(line.strip())
        
        docx_doc.save(output_path)
        logger.info(f"Filled PDF template converted to DOCX: {output_path}")
        
        return output_path
    
    def _format_citations(self, citations: List) -> str:
        """Format citations for display."""
        if not citations:
            return ""
        
        citation_texts = []
        for citation in citations[:3]:  # Limit to 3 citations
            parts = []
            # Handle both Citation objects and dicts
            if isinstance(citation, dict):
                doc_id = citation.get("doc_id")
                page_span = citation.get("page_span")
                citation_text = citation.get("citation_text")
            else:
                doc_id = getattr(citation, 'doc_id', None)
                page_span = getattr(citation, 'page_span', None)
                citation_text = getattr(citation, 'citation_text', None)
            
            if citation_text:
                citation_texts.append(citation_text)
            else:
                if doc_id:
                    parts.append(f"Doc: {doc_id}")
                if page_span:
                    if isinstance(page_span, (list, tuple)) and len(page_span) >= 2:
                        if page_span[0] == page_span[1]:
                            parts.append(f"P{page_span[0]}")
                        else:
                            parts.append(f"P{page_span[0]}-{page_span[1]}")
                if parts:
                    citation_texts.append(" | ".join(parts))
        
        return "; ".join(citation_texts) if citation_texts else ""

