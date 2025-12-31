"""DOCX parser for ParseForge."""

import logging
from pathlib import Path
from typing import List

from docx import Document as DocxDocument

from src.schema.document import BlockType, Document, Page, TextBlock, TitleBlock
from src.utils.exceptions import ParserError

logger = logging.getLogger(__name__)


def parse_docx(file_path: str) -> Document:
    """
    Parse DOCX file.

    Args:
        file_path: Path to DOCX file

    Returns:
        Document object
    """
    try:
        docx_doc = DocxDocument(file_path)
        blocks = []

        for para in docx_doc.paragraphs:
            if not para.text.strip():
                continue

            # Determine block type based on style
            style = para.style.name.lower()
            if "heading" in style or "title" in style:
                level = int(style[-1]) if style[-1].isdigit() else 1
                block = TitleBlock(
                    block_type=BlockType.TITLE,
                    text=para.text,
                    page_index=0,
                    level=level,
                )
            else:
                block = TextBlock(
                    block_type=BlockType.TEXT,
                    text=para.text,
                    page_index=0,
                )
            blocks.append(block)

        # Create page
        page = Page(page_index=0, width=612, height=792, blocks=blocks)  # Standard letter size

        # Create document
        document = Document(
            file_path=file_path,
            file_name=Path(file_path).name,
            pages=[page],
            total_pages=1,
        )

        return document

    except Exception as e:
        raise ParserError(f"Failed to parse DOCX: {e}") from e

